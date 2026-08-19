import os
import re
from docx import Document
from backend.translation.indictrans import translate_batch

def parse_tagged_text(tagged_text):
    """
    Parses a translated text containing HTML-like formatting tags (<b>, <i>, <u>)
    and returns a list of dictionaries with text segments and their style states.
    """
    # Split the text by tags, capturing the tags themselves
    parts = re.split(r'(</?[biu]>)', tagged_text)
    runs_data = []
    active_styles = {'b': False, 'i': False, 'u': False}
    
    for part in parts:
        if not part:
            continue
        
        if part == '<b>':
            active_styles['b'] = True
        elif part == '</b>':
            active_styles['b'] = False
        elif part == '<i>':
            active_styles['i'] = True
        elif part == '</i>':
            active_styles['i'] = False
        elif part == '<u>':
            active_styles['u'] = True
        elif part == '</u>':
            active_styles['u'] = False
        else:
            # This is standard text
            runs_data.append({
                'text': part,
                'bold': active_styles['b'],
                'italic': active_styles['i'],
                'underline': active_styles['u']
            })
            
    return runs_data

def paragraph_to_tagged_text(p):
    """
    Wraps styled runs in a paragraph with HTML-like tags for translation.
    """
    tagged_parts = []
    for run in p.runs:
        text = run.text
        if not text:
            continue
        
        bold = run.bold is True
        italic = run.italic is True
        underline = run.underline is True
        
        wrapped_text = text
        if bold:
            wrapped_text = f"<b>{wrapped_text}</b>"
        if italic:
            wrapped_text = f"<i>{wrapped_text}</i>"
        if underline:
            wrapped_text = f"<u>{wrapped_text}</u>"
            
        tagged_parts.append(wrapped_text)
        
    return "".join(tagged_parts)

def has_formatting_variation(p):
    """
    Checks if a paragraph has multiple runs with different formatting.
    """
    if len(p.runs) <= 1:
        return False
        
    # Get the style of the first run as reference
    first_r = p.runs[0]
    first_fmt = (first_r.bold is True, first_r.italic is True, first_r.underline is True)
    
    for r in p.runs[1:]:
        fmt = (r.bold is True, r.italic is True, r.underline is True)
        if fmt != first_fmt:
            return True
            
    return False

def translate_docx(input_path, output_path, batch_size=None, force_cpu=True):
    """
    Reads a DOCX file from input_path, translates all English text inside
    paragraphs and tables to Marathi, and saves the translated DOCX to output_path.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found at: {input_path}")
        
    doc = Document(input_path)
    
    translatable_items = []
    
    # 1. Collect all non-empty paragraphs from the main document body
    for p in doc.paragraphs:
        if p.text.strip():
            translatable_items.append({
                'type': 'paragraph',
                'object': p
            })
            
    # 2. Collect all non-empty paragraphs from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        translatable_items.append({
                            'type': 'paragraph',
                            'object': p
                        })
                        
    if not translatable_items:
        print("[DocTranslate] Extracted 0 paragraphs")
        # No text in document, just save it as is
        doc.save(output_path)
        return
        
    print(f"[DocTranslate] Extracted {len(translatable_items)} paragraphs")
        
    # 3. Prepare texts for batch translation
    texts_to_translate = []
    for item in translatable_items:
        p = item['object']
        if has_formatting_variation(p):
            item['has_variation'] = True
            item['tagged_text'] = paragraph_to_tagged_text(p)
            texts_to_translate.append(item['tagged_text'])
        else:
            item['has_variation'] = False
            item['tagged_text'] = p.text
            texts_to_translate.append(p.text)
            
    # 4. Perform translation of all items in a single batch call (which chunks internally)
    print(f"[DocTranslate] Translating {len(texts_to_translate)} sentences")
    translated_texts = translate_batch(texts_to_translate, batch_size=batch_size, force_cpu=force_cpu)
    print("[DocTranslate] Translation completed")
    
    # 5. Re-apply translated text to the document elements
    print("[DocTranslate] Creating translated DOCX")
    for item, translated_text in zip(translatable_items, translated_texts):
        p = item['object']
        
        # Save original font name and size from the first run (if any)
        orig_font_name = None
        orig_font_size = None
        orig_font_color = None
        orig_bold = False
        orig_italic = False
        orig_underline = False
        
        if p.runs:
            first_run = p.runs[0]
            orig_font_name = first_run.font.name
            orig_font_size = first_run.font.size
            if first_run.font.color and first_run.font.color.rgb:
                orig_font_color = first_run.font.color.rgb
            orig_bold = first_run.bold is True
            orig_italic = first_run.italic is True
            orig_underline = first_run.underline is True
            
        # Rebuild paragraph runs based on format variation
        if item['has_variation']:
            runs_data = parse_tagged_text(translated_text)
            p.text = "" # Clear original runs
            
            for r_data in runs_data:
                # Add a run with the text fragment
                new_run = p.add_run(r_data['text'])
                # Set formatting if it was active
                if r_data['bold']:
                    new_run.bold = True
                if r_data['italic']:
                    new_run.italic = True
                if r_data['underline']:
                    new_run.underline = True
                
                # Apply font properties from the original style
                if orig_font_name:
                    new_run.font.name = orig_font_name
                if orig_font_size:
                    new_run.font.size = orig_font_size
                if orig_font_color:
                    new_run.font.color.rgb = orig_font_color
        else:
            # Paragraph had uniform style, set translated text directly and apply formatting
            p.text = ""
            new_run = p.add_run(translated_text)
            if orig_bold:
                new_run.bold = True
            if orig_italic:
                new_run.italic = True
            if orig_underline:
                new_run.underline = True
                
            # Apply font properties
            if orig_font_name:
                new_run.font.name = orig_font_name
            if orig_font_size:
                new_run.font.size = orig_font_size
            if orig_font_color:
                new_run.font.color.rgb = orig_font_color
                
    doc.save(output_path)
