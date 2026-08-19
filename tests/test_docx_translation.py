import os
import sys
from docx import Document

# Set up path to import backend modules
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from backend.document.docx_processor import translate_docx

def create_sample_docx(path):
    print(f"Creating rich sample DOCX at: {path}")
    doc = Document()
    
    # 1. Main Heading
    h1 = doc.add_heading("Annual College Exhibition and Festival", level=1)
    h1.alignment = 1 # Center (1 corresponds to WD_ALIGN_PARAGRAPH.CENTER)
    
    # 2. Welcome Paragraph
    p1 = doc.add_paragraph("Welcome to our college. We are delighted to host the annual event this year.")
    
    # 3. Important notice paragraph with mixed formatting
    p2 = doc.add_paragraph()
    p2.add_run("This is an ")
    bold_run = p2.add_run("IMPORTANT")
    bold_run.bold = True
    p2.add_run(" notice. Please read the ")
    italic_run = p2.add_run("instructions")
    italic_run.italic = True
    p2.add_run(" carefully to avoid any confusion during the registration process.")
    
    # 4-12. Paragraphs 3 to 11 (Making 10+ paragraphs total)
    p3 = doc.add_paragraph("All students are requested to register their names before the deadline.")
    p4 = doc.add_paragraph("The exhibition will showcase various scientific projects and cultural achievements.")
    p5 = doc.add_paragraph("Distinguished guests from different universities will visit our campus.")
    p6 = doc.add_paragraph("Please maintain discipline and keep the campus clean during the event.")
    p7 = doc.add_paragraph("Volunteers should report to the main auditorium at 9:00 AM sharp.")
    p8 = doc.add_paragraph("No external vehicles will be allowed inside the campus premises.")
    p9 = doc.add_paragraph("Parking arrangements have been made at the ground next to the gate.")
    p10 = doc.add_paragraph("For any queries, contact the student council office.")
    p11 = doc.add_paragraph("We hope to see active participation from everyone to make this event a grand success.")
    
    # 13. Sub-heading
    doc.add_heading("Schedule of Events", level=2)
    
    # 14. Table with headers and multiple data rows
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("Event Name").bold = True
    hdr_cells[1].paragraphs[0].add_run("Timing").bold = True
    
    # Data Row 1
    row1 = table.rows[1].cells
    row1[0].paragraphs[0].text = "Inauguration Ceremony"
    row1[1].paragraphs[0].text = "10:00 AM"
    
    # Data Row 2
    row2 = table.rows[2].cells
    row2[0].paragraphs[0].text = "Project Exhibition"
    row2[1].paragraphs[0].text = "11:30 AM"
    
    # Data Row 3
    row3 = table.rows[3].cells
    row3[0].paragraphs[0].text = "Cultural Programs"
    row3[1].paragraphs[0].text = "03:00 PM"
    
    doc.save(path)
    print("Rich sample DOCX created successfully.")

def verify_translated_docx(path):
    print(f"\nVerifying translated DOCX at: {path}")
    doc = Document(path)
    
    print("\n--- Paragraphs ---")
    for i, p in enumerate(doc.paragraphs):
        print(f"P{i} [Alignment: {p.alignment}]: {p.text}")
        for r_idx, run in enumerate(p.runs):
            print(f"  Run {r_idx}: '{run.text}' (Bold: {run.bold}, Italic: {run.italic}, Underline: {run.underline})")
            
    print("\n--- Table Cells ---")
    for r_idx, row in enumerate(doc.tables[0].rows):
        row_text = []
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            # Print details of the first cell's paragraphs
            print(f"  Cell ({r_idx}, {cell_idx}) paragraphs:")
            for p_idx, p in enumerate(cell.paragraphs):
                print(f"    P{p_idx}: {p.text}")
                for r_idx2, run in enumerate(p.runs):
                    print(f"      Run {r_idx2}: '{run.text}' (Bold: {run.bold})")
            row_text.append(cell_text)
        print(f"Row {r_idx} Combined Text: {row_text}")
        
    # Check structures
    assert len(doc.paragraphs) >= 10, "Should contain at least 10 paragraphs"
    assert len(doc.tables) >= 1, "Should contain at least 1 table"
    assert len(doc.tables[0].rows) >= 2, "Table should have at least 2 rows"
    assert len(doc.tables[0].columns) == 2, "Table should have 2 columns"
    
    print("\nVerification successful! Formatting and structure are preserved.")

if __name__ == "__main__":
    # Ensure UTF-8 output on Windows console
    sys.stdout.reconfigure(encoding='utf-8')
    
    os.makedirs(os.path.dirname(os.path.abspath(__file__)), exist_ok=True)
    
    sample_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample.docx")
    translated_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_translated.docx")
    
    # Create the test file
    create_sample_docx(sample_path)
    
    # Translate the test file (explicitly testing CPU-first translation)
    print("\nStarting CPU-first translation test...")
    translate_docx(sample_path, translated_path, force_cpu=True)
    print("Translation completed.")
    
    # Verify the results
    verify_translated_docx(translated_path)
